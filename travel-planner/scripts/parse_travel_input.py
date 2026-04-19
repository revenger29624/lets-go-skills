#!/usr/bin/env python3
"""
旅行输入文件解析器
支持读取 .txt、.html、.docx 文件，提取纯文本内容用于旅行规划
"""
import sys
import re

def read_txt(file_path):
    """读取纯文本文件"""
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()

def read_html(file_path):
    """读取 HTML 文件，提取可读文本"""
    with open(file_path, 'r', encoding='utf-8') as f:
        html_content = f.read()
    
    # 移除 script 和 style 标签及其内容
    html_content = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
    html_content = re.sub(r'<style[^>]*>.*?</style>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
    
    # 将 HTML 标签替换为换行符
    html_content = re.sub(r'<br\s*/?>', '\n', html_content, flags=re.IGNORECASE)
    html_content = re.sub(r'<p[^>]*>', '\n', html_content, flags=re.IGNORECASE)
    html_content = re.sub(r'</p>', '', html_content, flags=re.IGNORECASE)
    html_content = re.sub(r'<div[^>]*>', '\n', html_content, flags=re.IGNORECASE)
    html_content = re.sub(r'</div>', '', html_content, flags=re.IGNORECASE)
    html_content = re.sub(r'<li[^>]*>', '\n- ', html_content, flags=re.IGNORECASE)
    html_content = re.sub(r'</li>', '', html_content, flags=re.IGNORECASE)
    
    # 移除剩余 HTML 标签
    html_content = re.sub(r'<[^>]+>', '', html_content)
    
    # 清理多余空白
    lines = [line.strip() for line in html_content.split('\n')]
    lines = [line for line in lines if line]
    
    return '\n'.join(lines)

def read_docx(file_path):
    """读取 Word 文档 (.docx)"""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # 读取表格
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(' | '.join(cells))
        
        return '\n'.join(paragraphs)
    except ImportError:
        return "[错误] 需要安装 python-docx 库来读取 Word 文档。运行: pip install python-docx"

def extract_travel_info(text):
    """从文本中提取旅行相关信息"""
    info = {}
    
    # 提取目的地
    destinations = re.findall(r'去([^\s,，。！？到]+?)(?:旅行|游玩|玩|天|晚|住)', text)
    if destinations:
        info['目的地'] = destinations[0]
    
    # 提取天数
    days = re.findall(r'(\d+)天(\d+)?晚?', text)
    if days:
        info['天数'] = days[0][0]
        info['晚数'] = days[0][1] if days[0][1] else str(int(days[0][0]) - 1)
    
    # 提取人数
    people = re.findall(r'(\d+)(?:人|位|个)?(?:朋友|人|情侣|夫妻|亲子|家人)', text)
    if people:
        info['人数'] = people[0]
    
    # 提取预算
    budget = re.findall(r'(\d+)(?:千|万)?元', text)
    if budget:
        info['预算'] = budget[0]
    
    return info

def main():
    if len(sys.argv) < 2:
        print("用法: python parse_travel_input.py <文件路径>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    ext = file_path.lower().split('.')[-1]
    
    if ext == 'txt':
        content = read_txt(file_path)
    elif ext == 'html' or ext == 'htm':
        content = read_html(file_path)
    elif ext == 'docx':
        content = read_docx(file_path)
    else:
        print(f"[错误] 不支持的文件格式: .{ext}")
        sys.exit(1)
    
    print("=== 文件内容 ===")
    print(content)
    
    # 尝试提取旅行信息
    info = extract_travel_info(content)
    if info:
        print("\n=== 识别的旅行信息 ===")
        for key, value in info.items():
            print(f"{key}: {value}")

if __name__ == '__main__':
    main()
